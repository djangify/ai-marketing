# content_generation/utils.py
import io
from docx import Document
from docx.shared import Pt
from django.utils.text import slugify

def generate_docx(title, content):
    """
    Generate a DOCX document from content
    
    Args:
        title (str): Title for the document
        content (str): Content text
        
    Returns:
        BytesIO: Document as a byte stream
    """
    document = Document()
    
    # Add title
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Add content - handle line breaks
    for paragraph_text in content.split('\n\n'):
        if paragraph_text.strip():
            paragraph = document.add_paragraph()
            paragraph.add_run(paragraph_text.strip())
    
    # Save to memory stream
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
